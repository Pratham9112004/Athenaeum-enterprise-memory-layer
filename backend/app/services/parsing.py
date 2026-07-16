"""Parsing uploaded files into page-tagged text.

PDFs are parsed page-by-page (so chunks can cite a page number). DOCX and plain
text/markdown have no real pages, so they yield a single page with ``page=None``.
Heavy parsers are imported lazily.
"""

from __future__ import annotations

from dataclasses import dataclass

from app.core.exceptions import ValidationError


@dataclass
class ParsedPage:
    page: int | None
    text: str


@dataclass
class ParsedDocument:
    pages: list[ParsedPage]

    @property
    def is_empty(self) -> bool:
        return not any(p.text.strip() for p in self.pages)


def parse_document(storage_path: str, extension: str) -> ParsedDocument:
    ext = extension.lower().lstrip(".")
    if ext == "pdf":
        return _parse_pdf(storage_path)
    if ext == "docx":
        return _parse_docx(storage_path)
    if ext in ("txt", "md"):
        return _parse_text(storage_path)
    raise ValidationError(f"Cannot parse unsupported extension '{extension}'.")


def _parse_pdf(path: str) -> ParsedDocument:
    import fitz  # PyMuPDF, lazy

    pages: list[ParsedPage] = []
    with fitz.open(path) as doc:
        for index, page in enumerate(doc):
            pages.append(ParsedPage(page=index + 1, text=page.get_text("text")))
    return ParsedDocument(pages=pages)


def _parse_docx(path: str) -> ParsedDocument:
    import docx  # python-docx, lazy

    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # Include table cell text, which python-docx keeps outside `paragraphs`.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return ParsedDocument(pages=[ParsedPage(page=None, text="\n".join(parts))])


def _parse_text(path: str) -> ParsedDocument:
    with open(path, encoding="utf-8", errors="replace") as fh:
        return ParsedDocument(pages=[ParsedPage(page=None, text=fh.read())])
