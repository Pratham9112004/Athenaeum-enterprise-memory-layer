"""Parsing tests using real PyMuPDF / python-docx on generated fixtures."""

import docx
import fitz  # PyMuPDF
import pytest

from app.core.exceptions import ValidationError
from app.services.parsing import parse_document


def test_parse_txt(tmp_path):
    path = tmp_path / "note.txt"
    path.write_text("The mitochondria is the powerhouse of the cell.", encoding="utf-8")
    parsed = parse_document(str(path), "txt")
    assert not parsed.is_empty
    assert "powerhouse" in parsed.pages[0].text
    assert parsed.pages[0].page is None


def test_parse_markdown(tmp_path):
    path = tmp_path / "doc.md"
    path.write_text("# Title\n\nSome **bold** body text.", encoding="utf-8")
    parsed = parse_document(str(path), "md")
    assert "body text" in parsed.pages[0].text


def test_parse_docx(tmp_path):
    path = tmp_path / "report.docx"
    document = docx.Document()
    document.add_paragraph("Quarterly revenue grew by twelve percent.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Region"
    table.rows[0].cells[1].text = "APAC"
    document.save(str(path))

    parsed = parse_document(str(path), "docx")
    text = parsed.pages[0].text
    assert "Quarterly revenue" in text
    assert "APAC" in text  # table cells are captured


def test_parse_pdf_is_page_tagged(tmp_path):
    path = tmp_path / "book.pdf"
    doc = fitz.open()
    for n in (1, 2):
        page = doc.new_page()
        page.insert_text((72, 72), f"This is page number {n}.")
    doc.save(str(path))
    doc.close()

    parsed = parse_document(str(path), "pdf")
    assert len(parsed.pages) == 2
    assert parsed.pages[0].page == 1
    assert parsed.pages[1].page == 2
    assert "page number 1" in parsed.pages[0].text


def test_empty_file_is_detected(tmp_path):
    path = tmp_path / "empty.txt"
    path.write_text("   \n  ", encoding="utf-8")
    assert parse_document(str(path), "txt").is_empty


def test_unsupported_extension_raises():
    with pytest.raises(ValidationError):
        parse_document("whatever.xyz", "xyz")
